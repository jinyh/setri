from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hs.font.size = Pt(22)
    elif level == 2:
        hs.font.size = Pt(16)
    else:
        hs.font.size = Pt(14)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


def add_title_page(doc):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('智能评审里程碑计划')
    run.font.size = Pt(28)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('（提纲模板）')
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for _ in range(3):
        doc.add_paragraph()

    info_items = [
        ('项目名称', 'Setri — 配电网业扩智能评审副驾驶（统一交付演示系统）'),
        ('文档版本', 'V0.3'),
        ('编制日期', '2026-04-15'),
        ('编制单位', '[编制单位名称]'),
        ('密级', '[内部/公开]'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(14)
        run.font.name = '仿宋'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    doc.add_page_break()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9E2F3',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elm)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


add_title_page(doc)

doc.add_heading('修订记录', level=1)
add_table(doc,
    ['版本', '日期', '修订内容', '修订人', '审核人'],
    [
        ['V0.3', '2026-04-15', '对齐 Milestone 计划 v0.3，新增里程碑详细计划、合同承接映射、风险管理与验收标准章节', '[修订人]', '[审核人]'],
        ['V0.1', '2026-04-10', '初版框架', '[修订人]', '[审核人]'],
    ]
)

doc.add_heading('目录', level=1)
toc_items = [
    '一、项目概述',
    '二、系统定位与目标',
    '三、总体架构与模块划分',
    '四、里程碑总览',
    '五、里程碑详细计划',
    '六、合同承接与成果映射',
    '七、关键技术路线',
    '八、资源与组织保障',
    '九、风险管理',
    '十、验收标准与交付物',
    '附录',
]
for item in toc_items:
    add_body(doc, item)
doc.add_page_break()

doc.add_heading('一、项目概述', level=1)

doc.add_heading('1.1 项目背景', level=2)
add_body(doc, '本项目面向上海经研院配电网评审专家，构建配电网业扩智能评审辅助系统（Setri）。系统以 10kV 开关站为切入点，通过"规范库→文件整理→专家知识库→交叉提炼→自动审查"五站数据管道，实现从原始资料到结构化审查意见的端到端智能化辅助评审能力。')
add_body(doc, '当前阶段重点为统一交付演示系统的构建（PRD v1.0），以统一平台形式承接 8 份合同项目材料的研究成果，形成可展示、可讲解、可验收的统一工作台。')

doc.add_heading('1.2 项目范围', level=2)
add_bullet(doc, '业务范围：10kV 配电网业扩项目评审，覆盖开关站、街坊站、电缆、通信等分项工程')
add_bullet(doc, '用户范围：业扩项目评审专家、技经评审人员、开关站专项评审人员、竣工验收与资料核查人员、项目管控人员')
add_bullet(doc, '资料范围：设计文件（送审版/审定版）、过程管理文件、技术规范文件、专家意见文件')

doc.add_heading('1.3 编制依据', level=2)
add_table(doc,
    ['序号', '依据文档', '说明'],
    [
        ['1', '《统一交付演示系统需求文档》', '统一需求基线'],
        ['2', '《产品需求文档 PRD v1.0》', '现行主 PRD'],
        ['3', '《统一交付演示系统重构计划》', '工程重构与首期边界'],
        ['4', '《统一交付演示系统 Milestone 计划》', '阶段交付顺序'],
        ['5', '《合同承接矩阵》', '8 份合同成果到系统模块的追踪'],
        ['6', '《多模态审核技术路线说明》', 'P3-P5 方法层设计'],
        ['7', '《专家知识库与规则演化设计》', '知识库与规则演化方法'],
        ['8', '《评测与执行底座工程设计》', '评测、执行 trace 与发布门槛'],
    ]
)

doc.add_page_break()
doc.add_heading('二、系统定位与目标', level=1)

doc.add_heading('2.1 系统定位', level=2)
add_body(doc, '面向经研院专家与业务人员的配电网项目智能审查/核查统一工作台。不是多套独立系统的拼接，也不是纯算法验证台，而是一套可基于真实资料包完成"导入、解析、核查、追证、输出"的统一业务型演示系统。')

doc.add_heading('2.2 核心目标', level=2)
add_bullet(doc, '可讲述：甲方能清晰理解统一平台与各专题成果的关系')
add_bullet(doc, '可操作：用户能在浏览器内完成完整演示流程')
add_bullet(doc, '可扩展：后续合同项和专题能力可持续挂接到统一架构')
add_bullet(doc, '可验收：系统能够说明每项研究成果在平台中的承载位置')

doc.add_heading('2.3 非目标', level=2)
add_bullet(doc, '不按"每份合同一套独立系统"方式建设')
add_bullet(doc, '不以"全自动替代专家"为目标')
add_bullet(doc, '不在首期一次性完成全部 P3-P5 真能力')
add_bullet(doc, '不覆盖全部电压等级和全部工程类型')

doc.add_page_break()
doc.add_heading('三、总体架构与模块划分', level=1)

doc.add_heading('3.1 系统总体结构', level=2)
add_body(doc, '统一交付演示系统采用"公共能力底座 + 核心主线模块 + 专题扩展模块"的结构：')

add_table(doc,
    ['层级', '模块', '定位', '首期策略'],
    [
        ['公共底座', '项目与案例管理', '项目创建、资料导入、状态管理', '必须完成'],
        ['公共底座', '多模态资料接入与解析', 'PDF/Word/图纸/表格等资料输入与结构化', '必须完成'],
        ['公共底座', '知识与规则中心', '规范条款、业务规则、专家知识统一管理', '基础闭环'],
        ['公共底座', '核查与推理引擎', '一致性校验、规则匹配、问题定位', '基础闭环'],
        ['公共底座', 'AI 助手与问答', '基于知识库的问答与检索', '入口预留'],
        ['公共底座', '结果输出与证据链展示', '问题清单、证据链、结论摘要', '必须完成'],
        ['核心主线 M1', '业扩项目智能评审', '统一评审主线入口', '首期真实主线骨架'],
        ['核心主线 M2', '技经资料预审与评审', '技经资料解析、校验与问题定位', '首期真实闭环'],
        ['专题扩展 M3', '开关站专家知识库', '专家经验、规则条件、知识查询与演化', '首期保留骨架'],
        ['专题扩展 M4', '工程智慧进度管控', '进度感知、风险预警、纠偏仿真', '首期保留骨架'],
        ['专题扩展 M5', '竣工资料智能核查', '图纸-实物-档案关联与三类核查', '首期保留骨架'],
    ]
)

doc.add_heading('3.2 数据管道 P1-P5 与产品模块的关系', level=2)
add_table(doc,
    ['能力链路', '作用', '首期要求'],
    [
        ['P1 规范库构建', '规范源登记、扫描、摘要、产物管理', '真实闭环'],
        ['P2 文件整理与解析', '项目资料归档、路径映射、基础解析', '真实闭环'],
        ['P3 专家知识库', '专家经验、知识抽取、更新演化', '产品骨架'],
        ['P4 交叉提炼', '规则形成、规则维护、规则评估', '产品骨架'],
        ['P5 自动审查', '规则执行、结果生成、意见管理', '产品骨架'],
    ]
)

doc.add_heading('3.3 统一业务主线', level=2)
add_body(doc, '所有模块复用同一条业务主线：')
steps = [
    '选择项目或案例',
    '导入资料或过程数据',
    '进行解析、抽取和结构化',
    '调用知识库、规则库或图谱进行匹配',
    '执行核查、推理或预警分析',
    '展示问题、证据与建议',
    '生成人工可确认的最终结论或报告',
]
for i, s in enumerate(steps, 1):
    add_body(doc, f'  {i}. {s}')

doc.add_heading('3.4 技术栈', level=2)
add_table(doc,
    ['层级', '技术选型'],
    [
        ['前端', 'React + Next.js + TypeScript + shadcn/ui + Tailwind CSS'],
        ['后端 API', 'FastAPI'],
        ['后台任务', 'Celery + Redis'],
        ['持久化', '文件系统 + JSON 元数据优先，PostgreSQL 预留'],
        ['VLM', 'Gemini 3.0 Flash（MVP）→ Qwen VL（微调私有化）'],
        ['认证', 'JWT（内网部署）'],
        ['部署', 'Docker 化（后续交付增强项）'],
    ]
)

doc.add_page_break()
doc.add_heading('四、里程碑总览', level=1)

doc.add_heading('4.1 里程碑阶段划分', level=2)
add_table(doc,
    ['Milestone', '名称', '目标', '建议周期', '主要产出'],
    [
        ['M1', '需求冻结与案例范围确认', '统一需求基线、确认模块边界与演示案例', '第 1-2 周', '需求文档、能力映射表、演示案例清单'],
        ['M2', '统一平台骨架与信息架构确定', '建立信息架构和页面主骨架', '第 3-5 周', '模块结构、页面流转、演示骨架'],
        ['M3', '核心主线能力完成', '跑通业扩评审与技经评审核心主线', '第 6-10 周', '核心主线演示版本、核心案例结果页'],
        ['M4', '专题能力完成', '挂载开关站知识库、进度管控、竣工核查专题', '第 11-15 周', '专题演示版本、专题案例结果页'],
        ['M5', '联调、验收与演示材料收口', '形成可对内汇报、对外验收的统一版本', '第 16-18 周', '联调版系统、演示手册、验收说明'],
    ]
)

doc.add_heading('4.2 工程实施阶段（重构计划）', level=2)
add_table(doc,
    ['阶段', '名称', '内容', '完成标志'],
    [
        ['A', '基线冻结与清场准备', '盘点旧资产、形成映射说明、明确删除清单', '旧资产清单确认'],
        ['B', '新骨架建立', '初始化 monorepo（apps/packages/tests）', '新目录结构可运行'],
        ['C', '领域层重建', '重写 P1/P2 schema、迁移同义词映射', '核心领域模型通过单测'],
        ['D', 'P1 接入', '收编扫描脚本、API-worker-结果闭环', 'P1 任务可创建执行查询'],
        ['E', 'P2 接入', '中文标准目录整理流程、映射记录闭环', 'P2 任务可创建执行查询'],
        ['F', '统一工作台接入', '首页/导航/任务面板、P3-P5 骨架页', '前端可展示真任务'],
        ['G', '旧主线退出', '删除旧目录、更新文档', '不依赖旧 CLI/旧 web'],
        ['H', 'P3 专家知识库接入', '后续迭代', 'P2 结构化结果稳定后触发'],
        ['I', 'P4 交叉提炼接入', '后续迭代', 'P1+P3 具备稳定输出后触发'],
        ['J', 'P5 自动审查接入', '后续迭代', 'P4 可执行规则可用后触发'],
    ]
)

doc.add_heading('4.3 推荐交付优先级', level=2)
priorities = [
    '统一需求与案例范围',
    '统一平台骨架',
    '业扩评审主线（M1）',
    '技经评审主线（M2）',
    '开关站知识库专题（M3）',
    '竣工核查专题（M5）',
    '进度管控专题（M4）',
    '联调与验收材料',
]
for i, p in enumerate(priorities, 1):
    add_body(doc, f'  {i}. {p}')

doc.add_page_break()
doc.add_heading('五、里程碑详细计划', level=1)

doc.add_heading('5.1 M1：需求冻结与案例范围确认', level=2)
add_table(doc,
    ['项目', '内容'],
    [
        ['目标', '将 8 份项目材料统一归并为一个产品需求基线'],
        ['重点工作', '完成统一需求文档；建立能力与系统模块映射表；确认首版业务闭环；确认每个模块代表性案例'],
        ['完成标准', '需求按系统能力组织而非按材料编号；模块边界明确；案例清单可支撑后续验证'],
        ['依赖', '合同目录材料完整可读；业务侧认可统一平台定位'],
        ['交付物', '需求文档、能力映射表、案例清单'],
    ]
)

doc.add_heading('5.2 M2：统一平台骨架与信息架构确定', level=2)
add_table(doc,
    ['项目', '内容'],
    [
        ['目标', '建立统一演示系统的信息架构和页面主骨架'],
        ['重点工作', '定义首页/案例入口/模块入口/结果页结构；确定统一主线；明确公共能力底座承载方式'],
        ['完成标准', '各模块共享同一套主线；首页区分核心主线与专题入口；输入/处理/输出位置明确'],
        ['依赖', 'M1 需求边界和案例范围已冻结'],
        ['交付物', '平台信息架构、页面清单、流程图、模块边界说明'],
    ]
)

doc.add_heading('5.3 M3：核心主线能力完成', level=2)
add_table(doc,
    ['项目', '内容'],
    [
        ['目标', '跑通最能代表统一平台价值的两条核心主线'],
        ['范围', 'M1 业扩项目智能评审 + M2 技经资料预审与评审'],
        ['重点工作', '打通资料导入/解析/规则匹配/问题输出；打通技经字段抽取/逻辑校验/建议输出；M2 至少基于可研/初设报告、概算/估算书、图纸/材料清册三类稳定资料完成结果展示，并支持至少 1 组跨文档一致性校验；建立统一证据链展示'],
        ['完成标准', '端到端主线可演示；结论可回溯资料和规则；输出可用于评审讲解'],
        ['依赖', 'M1 平台骨架已具备；首批案例数据可用'],
        ['交付物', '核心主线演示版本、核心案例结果页、主线讲解脚本'],
    ]
)

doc.add_heading('5.4 M4：专题能力完成', level=2)
add_table(doc,
    ['项目', '内容'],
    [
        ['目标', '在统一平台上挂载剩余专题能力，体现项目群完整性'],
        ['范围', 'M3 开关站专家知识库 + M4 工程智慧进度管控 + M5 竣工资料智能核查'],
        ['重点工作', '专家知识展示/查询/问答/更新记录演示；进度预警/风险证据链/纠偏仿真演示；竣工图谱/三类核查/问题定位演示'],
        ['完成标准', '三个专题可独立演示；复用公共证据与结果表达；讲解时可说明与核心主线关系'],
        ['依赖', 'M3 公共能力和结果框架已稳定'],
        ['交付物', '专题演示版本、专题案例结果页、专题讲解脚本'],
    ]
)

doc.add_heading('5.5 M5：联调、验收与演示材料收口', level=2)
add_table(doc,
    ['项目', '内容'],
    [
        ['目标', '形成可对内汇报、对外验收的统一演示版本'],
        ['重点工作', '全模块联调与修正；固化演示脚本/讲解顺序/亮点话术；准备验收口径/案例说明/结果截图'],
        ['完成标准', '至少 1 条核心主线和 3 个专题模块可稳定演示；演示前后连贯；可说明每项研究成果承载位置'],
        ['依赖', 'M3、M4 均完成并通过内部评审'],
        ['交付物', '联调版演示系统、演示手册、验收说明、答疑清单'],
    ]
)

doc.add_page_break()
doc.add_heading('六、合同承接与成果映射', level=1)

doc.add_heading('6.1 合同材料总览', level=2)
add_table(doc,
    ['序号', '合同材料名称', '状态', '主要模块', '首期承接方式'],
    [
        ['1', '2024 配电网业扩项目智能化审查与合规性评估', '已验收', 'M1', '统一评审主线、规则中心、结果追溯表达'],
        ['2', '2025 基于电力行业大模型的技经资料智能辅助预审', '执行中', 'M2', 'P2 真实闭环 + 技经解析/校验主线'],
        ['3', '2025 基于知识图谱的评审规则自动化匹配研究', '已验收', 'M1', '规则来源、规则命中、证据链表达'],
        ['4', '2025 基于知识图谱的技经智能辅助评审与优化', '执行中', 'M2', '技经逻辑校验与问题输出主线'],
        ['5', '2025 开关站专家知识库构建辅助评审方法研究', '已验收', 'M3', '专题骨架、知识承接说明、入口预留'],
        ['6', '2026 基于智能体协同的智能评审关键技术研究', '执行中', 'M1', '多智能体入口、任务拆解、问答扩展位预留'],
        ['7', '2026 工程智慧进度管控方法与验证', '招标中', 'M4', '专题骨架、数据对象预留、结果表达预留'],
        ['8', '2026 竣工资料多模态理解与智能核查方法研究', '招标中', 'M5', '专题骨架、图谱/核查对象预留、证据链预留'],
    ]
)

doc.add_heading('6.2 承接原则', level=2)
add_bullet(doc, '已验收合同：在统一系统中体现成果承接位置和复用价值，不要求首期全部重做为真能力')
add_bullet(doc, '执行中合同：必须明确模块落位、交付阶段和后续实现抓手，不能遗漏')
add_bullet(doc, '招标中合同：必须在统一系统结构中预留专题入口、对象边界和结果表达方式')
add_bullet(doc, '论文、专利、研究报告等成果不直接转译为独立页面，仅作为系统能力来源说明保留')

doc.add_page_break()
doc.add_heading('七、关键技术路线', level=1)

doc.add_heading('7.1 多模态审核技术路线', level=2)
add_bullet(doc, 'PDF/Word/扫描件/图纸/表格/图片等多格式资料输入支持')
add_bullet(doc, '关键字段抽取、表格恢复、版面理解、印章/签字识别')
add_bullet(doc, '输出结构化字段及置信度')
add_bullet(doc, 'LLM 不负责算术计算、阈值比较或确定性规则执行')
add_bullet(doc, '关键数字抽取结果保留置信度或来源追踪信息')

doc.add_heading('7.2 专家知识库与规则演化', level=2)
add_bullet(doc, '三元关联：专家意见 ↔ 设计文件 ↔ 引用规范')
add_bullet(doc, 'P5 主评审链路：规则库直接评审设计文件')
add_bullet(doc, '专家意见匹配用于衡量意见与规则体系的一致性，不替代主评审链路')
add_bullet(doc, '专家反馈式微调：对规则命中、表达、适用范围和误报漏报的修正')
add_bullet(doc, '规则库评审能力演化以实际表现为主观测指标')

doc.add_heading('7.3 评测与执行底座', level=2)
add_bullet(doc, '执行 trace 记录：输入来源、字段来源、规则命中来源、输出结论来源')
add_bullet(doc, '评测矩阵与评分方式')
add_bullet(doc, '回归门槛与发布准入机制')
add_bullet(doc, '所有自动结论可回溯到资料片段、规则条款、知识节点或历史案例之一')

doc.add_page_break()
doc.add_heading('八、资源与组织保障', level=1)

doc.add_heading('8.1 团队组织（模板）', level=2)
add_table(doc,
    ['角色', '职责', '人员', '备注'],
    [
        ['项目经理', '项目整体管理与协调', '[待填]', ''],
        ['产品负责人', '需求管理与产品决策', '[待填]', ''],
        ['后端开发', 'API/Worker/Pipeline 开发', '[待填]', ''],
        ['前端开发', '统一工作台 UI 开发', '[待填]', ''],
        ['AI/算法', 'VLM/LLM 集成与优化', '[待填]', ''],
        ['测试', '质量保障与验收测试', '[待填]', ''],
    ]
)

doc.add_heading('8.2 环境与工具', level=2)
add_table(doc,
    ['项目', '说明'],
    [
        ['开发环境', 'Python 3.10+，Node.js 18+，uv 包管理'],
        ['版本管理', 'Git monorepo'],
        ['CI/CD', '[待确定]'],
        ['部署环境', '内网本地服务器，Docker 化'],
        ['外部服务', 'Gemini 3.0 Flash / Qwen VL（VLM），Redis'],
    ]
)

doc.add_page_break()
doc.add_heading('九、风险管理', level=1)

add_table(doc,
    ['编号', '风险描述', '影响', '概率', '应对措施'],
    [
        ['R1', '模块过多导致系统像"拼盘"', '高', '中', '始终以统一主线为骨架；共享结果表达方式；首页只突出核心主线与专题扩展'],
        ['R2', '真实资料质量不稳定', '中', '高', '每个模块至少准备 1 套稳定样例；对真实案例与演示样例分层管理'],
        ['R3', '专题能力深浅不一', '中', '高', '首版以"可演示"而非"全自动最优"为完成标准；共性能力优先做深'],
        ['R4', '验收时无法说明研究成果承接关系', '高', '低', '准备"项目成果映射表"；每个模块明确对应项目成果'],
        ['R5', '首期引入复杂基础设施拖慢进度', '中', '中', '文件系统 + JSON 元数据起步；数据库接口预留但不阻塞'],
        ['R6', '误把旧中间产物当成可持续骨架', '高', '中', '迁移知识资产不迁移旧骨架；先做保留/废弃分类再迁移'],
        ['R7', 'P1/P2 与演示工作台双线推进范围失控', '高', '中', '第一阶段只做 P1/P2 真能力；P3-P5 只保留骨架不提前实现'],
    ]
)

doc.add_page_break()
doc.add_heading('十、验收标准与交付物', level=1)

doc.add_heading('10.1 首期验收标准', level=2)
add_table(doc,
    ['编号', '验收项', '验收标准', '优先级'],
    [
        ['V1', '统一平台', '系统为统一平台而非散装拼接', 'P0'],
        ['V2', '统一业务主线', '系统有清晰的公共主线', 'P0'],
        ['V3', 'P1 规范库闭环', '可创建/执行任务、查询状态、查看产物', 'P0'],
        ['V4', 'P2 文件整理闭环', '可创建/执行任务、查询状态、查看路径映射', 'P0'],
        ['V5', '至少 1 个样例项目跑通', '真实项目数据可跑通新 P2 链路', 'P0'],
        ['V6', '前端可展示', '任务状态、结果摘要和产物入口可展示', 'P0'],
        ['V7', 'P3-P5 骨架页', '存在且明确标示阶段状态，不误导用户', 'P1'],
        ['V8', '旧主线退出', '新工程不依赖旧 CLI 入口和旧占位骨架', 'P1'],
        ['V9', '合同成果映射', '可说明每项研究成果在系统中的承接位置', 'P1'],
        ['V10', '前端/API/Worker 独立启动', '三个组件可独立启动运行', 'P0'],
    ]
)

doc.add_heading('10.2 交付物清单', level=2)
add_table(doc,
    ['Milestone', '交付物'],
    [
        ['M1', '需求文档、能力映射表、演示案例清单'],
        ['M2', '平台信息架构、页面清单、流程图、模块边界说明'],
        ['M3', '核心主线演示版本、核心案例结果页、主线讲解脚本'],
        ['M4', '专题演示版本、专题案例结果页、专题讲解脚本'],
        ['M5', '联调版演示系统、演示手册、验收说明、答疑清单'],
    ]
)

doc.add_heading('10.3 验收方式', level=2)
add_bullet(doc, '功能演示：基于真实项目资料包的端到端演示')
add_bullet(doc, '文档审查：需求文档、设计文档、验收说明的完整性审查')
add_bullet(doc, '代码审查：核心模块代码质量与测试覆盖')
add_bullet(doc, '成果映射检查：合同承接矩阵逐条核对')

doc.add_page_break()
doc.add_heading('附录', level=1)

doc.add_heading('附录 A：首期页面范围', level=2)
add_table(doc,
    ['序号', '页面', '说明'],
    [
        ['1', '系统首页', '统一叙事，展示系统定位、模块、案例入口'],
        ['2', '模块导航页', '核心主线与专题扩展入口'],
        ['3', '项目/案例列表页', '演示项目管理'],
        ['4', '任务列表页', '统一任务中心，支持过滤与查询'],
        ['5', 'P1 工作区', '规范库构建任务创建与结果查看'],
        ['6', 'P2 工作区', '文件整理任务创建与结果查看'],
        ['7', '任务详情页', '状态、进度、结果摘要'],
        ['8', '规范库浏览页', '规范库结果列表与详情'],
        ['9', '结果详情页', '问题清单、证据链、结论'],
        ['10', 'M3-M5 专题骨架页', '保留入口与能力说明'],
    ]
)

doc.add_heading('附录 B：首期最小 API 端点', level=2)
add_table(doc,
    ['端点', '方法', '作用'],
    [
        ['/api/projects', 'GET', '项目/案例列表'],
        ['/api/projects/{project_id}', 'GET', '项目详情'],
        ['/api/pipelines/p1/tasks', 'POST', '创建 P1 任务'],
        ['/api/pipelines/p2/tasks', 'POST', '创建 P2 任务'],
        ['/api/tasks', 'GET', '任务列表（支持过滤）'],
        ['/api/tasks/{task_id}', 'GET', '任务详情'],
        ['/api/tasks/{task_id}/artifacts', 'GET', '任务产物列表'],
        ['/api/regulations', 'GET', '规范库结果列表'],
        ['/api/regulations/{result_id}', 'GET', '规范库结果详情'],
    ]
)

doc.add_heading('附录 C：工程目录结构', level=2)
add_body(doc, 'monorepo 结构如下：')
structure = """\
/
├── apps/
│   ├── web/          # Next.js 工作台
│   ├── api/          # FastAPI HTTP API
│   └── worker/       # Celery 后台任务执行器
├── packages/
│   ├── domain/       # 领域模型、schema、枚举
│   ├── pipelines/    # P1/P2 管道逻辑
│   └── connectors/   # 文件系统、PDF、LLM/VLM 适配层
├── tests/
│   ├── unit/
│   ├── integration/
│   ├── fixtures/
│   └── e2e/
├── data/             # 演示样例、测试夹具、产物
├── docs/
└── scripts/"""
for line in structure.split('\n'):
    add_body(doc, f'  {line}')

doc.add_heading('附录 D：关键术语表', level=2)
add_table(doc,
    ['术语', '含义'],
    [
        ['送审版', '专家收到的待审材料（修改前），系统输入基准'],
        ['审定版', '专家评审后的修正版本（修改后），参考标准答案'],
        ['三元关联', '专家意见 ↔ 设计文件 ↔ 引用规范的关联结构'],
        ['管道站', '数据管道中的一个处理阶段（P1-P5）'],
        ['交叉提炼', '"应然"（规范）× "实然"（专家）→ 可执行规则'],
        ['公共底座', '所有模块共享的基础能力层'],
        ['核心主线', '业扩评审（M1）+ 技经评审（M2）'],
        ['专题扩展', '开关站知识库（M3）+ 进度管控（M4）+ 竣工核查（M5）'],
    ]
)

output_path = r'D:\code_cys\stategrid_setri-main\setri-main\智能评审里程碑计划【提纲模板】.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
